import feedparser
import ssl
import re
from newspaper import Article
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime
import schedule
import time
import logging
from googlenewsdecoder import new_decoderv1

# Email configuration
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
EMAIL_ADDRESS = "jonmay002@gmail.com"
EMAIL_PASSWORD = "wioe hgue tlnu gzah"

# Logging configuration
logging.basicConfig(filename='failed_articles.log', level=logging.ERROR, format='%(asctime)s - %(message)s')

def log_failure(url, error_message):
    logging.error(f"URL: {url}\nError: {error_message}")

def generate_rss_url(entity, time_range):
    base_url = "https://news.google.com/rss/search?q="
    query = entity.replace(" ", "+") + f"+when:{time_range}"
    url = f"{base_url}{query}&ceid=US:en&hl=en-US&gl=US"
    return url

def fetch_articles(rss_url):
    if hasattr(ssl, '_create_unverified_context'):
        ssl._create_default_https_context = ssl._create_unverified_context
    feed = feedparser.parse(rss_url)
    articles = []
    for entry in feed.entries:
        articles.append({
            "title": entry.title,
            "link": entry.link,
            "published": entry.published,
            "source": entry.source.get("title", "Unknown Source") if hasattr(entry, 'source') else "Unknown Source"
        })
    return articles

def resolve_actual_url(rss_link, interval_time=1):
    try:
        decoded_url = new_decoderv1(rss_link, interval=interval_time)
        if decoded_url.get("status"):
            return decoded_url["decoded_url"]
        else:
            log_failure(rss_link, decoded_url.get("message", "Unknown decoding error"))
            return rss_link
    except Exception as e:
        log_failure(rss_link, str(e))
        return rss_link

def extract_article_content_with_retries(url, retries=3, delay=5):
    for attempt in range(retries):
        try:
            article = Article(url)
            article.download()
            article.parse()
            return {
                'title': article.title,
                'text': article.text if len(article.text) >= 200 else "Article text unavailable",
                'url': url,
                'published': article.publish_date.strftime('%Y-%m-%d') if article.publish_date else "Unknown"
            }
        except Exception as e:
            print(f"Attempt {attempt + 1} failed for {url}: {e}")
            time.sleep(delay)
    # Log failure after retries
    log_failure(url, f"Failed after {retries} retries.")
    return {
        'title': "Title Unavailable",
        'text': "Error extracting content after multiple attempts.",
        'url': url,
        'published': "Unknown"
    }



def process_text_manual(public_figure, text, url):
    if not text.strip() or text == "Error extracting content after multiple attempts.":
        return "Article text unavailable. See the link above to extract manually."
    
    # Generate search terms
    search_terms = [public_figure.lower()]  # Full name
    name_parts = public_figure.lower().split()  # First and last name
    name_parts = [name_parts[-1]]
    search_terms.extend(name_parts)  # Add individual parts

    # Sort terms by length (to prioritize longest matches first)
    search_terms = sorted(search_terms, key=len, reverse=True)

    # Prepare regex pattern for all terms (prioritize longest matches)
    search_pattern = r'\b(?:' + '|'.join(re.escape(term) for term in search_terms) + r')\b'

    # Prepare output
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]  # Remove empty lines
    relevant_paragraphs = []
    mention_count = 0

    for i, paragraph in enumerate(paragraphs):
        paragraph_lower = paragraph.lower()
        
        # Find all matches in the paragraph using regex
        matches = re.findall(search_pattern, paragraph_lower, flags=re.IGNORECASE)
        if matches:
            # Increment mention count by the number of unique matches
            mention_count += len(set(matches))
            
            # Add the paragraph before, the current paragraph, and the paragraph after
            if i > 0:  # One before
                relevant_paragraphs.append(paragraphs[i - 1])
            relevant_paragraphs.append(paragraph)  # Current
            if i < len(paragraphs) - 1:  # One after
                relevant_paragraphs.append(paragraphs[i + 1])

    if not relevant_paragraphs:
        return "No mentions identified. Reviewing for relevance is recommended."

    # Remove duplicates while preserving order
    seen = set()
    unique_paragraphs = []
    for paragraph in relevant_paragraphs:
        if paragraph not in seen:
            unique_paragraphs.append(paragraph)
            seen.add(paragraph)

    # Format mentions in all capitals
    formatted_paragraphs = []
    for paragraph in unique_paragraphs:
        formatted_paragraph = paragraph
        for term in search_terms:
            # Replace case-insensitively using regex
            formatted_paragraph = re.sub(rf'\b{re.escape(term)}\b', term.upper(), formatted_paragraph, flags=re.IGNORECASE)
        formatted_paragraphs.append(formatted_paragraph)

    # Combine results and append the total mentions count
    result = '\n\n'.join(formatted_paragraphs)
    result += f"\n\nTOTAL MENTIONS: {mention_count}"
    return result




def create_word_document(public_figure, articles):
    filename = f"{datetime.now().strftime('%Y-%m-%d')} {public_figure} Clips.docx"
    doc = Document()

    # Set document title
    title = f"{public_figure} Press Clips {datetime.now().strftime('%Y.%m.%d')}"
    title_paragraph = doc.add_heading(level=0)
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add articles
    for idx, article in enumerate(articles, start=1):
        subtitle = f"Article {idx}: {article['title']}"
        subtitle_paragraph = doc.add_heading(level=2)
        run = subtitle_paragraph.add_run(subtitle)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 139)

        publication_date = doc.add_paragraph()
        publication_date.add_run("Publication Date: ").bold = True
        run = publication_date.add_run(article['published'])
        run.font.name = 'Arial'

        url_paragraph = doc.add_paragraph()
        url_paragraph.add_run("Article URL: ").bold = True
        run = url_paragraph.add_run(article['url'])
        run.font.name = 'Arial'

        processed_text_paragraph = doc.add_paragraph()
        processed_text_paragraph.add_run("Processed Text:\n").bold = True
        run = processed_text_paragraph.add_run(article['processed_text'])
        run.font.name = 'Arial'

    # Apply default font to all paragraphs
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    doc.save(filename)
    return filename

def send_email(recipient, subject, body, attachment):
    msg = MIMEMultipart()
    msg['From'] = EMAIL_ADDRESS
    msg['To'] = recipient
    msg['Subject'] = subject
    with open(attachment, "rb") as file:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(file.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f"attachment; filename={attachment}")
    msg.attach(part)
    with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.send_message(msg)

def fetch_and_send(public_figure, recipient, time_range):
    rss_url = generate_rss_url(public_figure, time_range)
    articles = fetch_articles(rss_url)
    processed_articles = []
    for article in articles:
        article['url'] = resolve_actual_url(article['link'])
        content = extract_article_content_with_retries(article['url'])
        if content:
            article['title'] = content['title']
            article['published'] = content['published']
            article['processed_text'] = process_text_manual(public_figure, content['text'], article['url'])
            processed_articles.append(article)  # Only include successfully processed articles
        else:
            print(f"Skipping article due to extraction failure: {article['url']}")
    if processed_articles:
        filename = create_word_document(public_figure, processed_articles)
        send_email(recipient, f"{public_figure} Clips {datetime.now().strftime('%Y.%m.%d')}", "[AUTOMATED MESSAGE] Please find the updates attached.", filename)
    else:
        print("No articles were successfully processed.")

# Scheduling example
def schedule_emails():
    schedule.every().monday.at("10:17").do(fetch_and_send, "Justin Baldoni", "gjkim@college.harvard.edu", "24h")
    schedule.every().tuesday.at("14:39").do(fetch_and_send, "Ethan Mollick", "gjkim@college.harvard.edu", "7d")
    schedule.every().tuesday.at("06:30").do(fetch_and_send, "Jeff Hurd", "jonmay@college.harvard.edu", "24h")
    schedule.every().wednesday.at("06:30").do(fetch_and_send, "Jeff Hurd", "jonmay@college.harvard.edu", "24h")
    schedule.every().thursday.at("13:10").do(fetch_and_send, "John Doe", "jonmay@college.harvard.edu", "24h")
    schedule.every().friday.at("17:06").do(fetch_and_send, "Jeff Hurd", "jonmay@college.harvard.edu", "24h")
    schedule.every().sunday.at("20:34").do(fetch_and_send, "Justin Baldoni", "gjkim@college.harvard.edu", "24h")
    while True:
        schedule.run_pending()
        time.sleep(1)

if __name__ == "__main__":
    schedule_emails();